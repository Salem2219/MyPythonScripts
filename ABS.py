import re
import os
import shelve
import shutil
import pyperclip
import send2trash
import sys
import webbrowser
from selenium import webdriver
import openpyxl
import PyPDF2
import docx
import smtplib
import imapclient
import pyzmail
import pyautogui


# Regular Expressions
def find_phone_num(message):
    phone_num_regex = re.compile(r'(\d\d\d)-(\d\d\d-\d\d\d\d)')
    mo = phone_num_regex.findall(message)
    return mo


def remove_nonletter(title: str):
    regex = re.compile(r'\W')
    t = regex.findall(title)
    for i in t:
        if i != ' ':
            title = title.replace(i, '')
    return title


def xmas(text):
    xmas_regex = re.compile(r'\d+\s\w+')
    return xmas_regex.findall(text)


def vowels(message):
    vowel_regex = re.compile(r'[aeiou]', re.I)
    print(vowel_regex.findall(message))


def begins_with_hello(messge):
    begins_with_helo_regex = re.compile(r'^Hello')
    return begins_with_helo_regex.search(messge).group()


def end_with_world(message):
    end_with_world_regex = re.compile(r'world!$')
    return end_with_world_regex.search(message).group()


def allDigit(text):
    all_digits_regex = re.compile(r'^\d+$')
    return all_digits_regex.search(text).group()


def at(text):
    atRegex = re.compile(r'.{1,2}at')
    return atRegex.findall(text)


def name(text):
    nameRegex = re.compile('First Name: (.*) Last Name: (.*)')
    return nameRegex.findall(text)


def nongreedy(text):
    nongreedy_regex = re.compile(r'<(.*?)>')
    return nongreedy_regex.findall(text)


def dotStar(text):
    dot_star = re.compile(r'.*', re.DOTALL)
    return dot_star.search(text).group()


def names(text):
    names_regex = re.compile('Agent \w+')
    return names_regex.sub('REDACTED', text)


def names2(text):
    names_regex = re.compile('Agent (\w)\w*')
    return names_regex.sub('Agent \1****', text)


# Files and Folders

def fundamentals():
    print(os.path.join('folder1', 'folder2', 'folder3', 'file.png'))
    print(os.sep)
    os.chdir('c:\\')
    print(os.getcwd())
    print(os.path.abspath('spam.png'))
    print(os.path.abspath('..\\..\\spam.png'))
    print(os.path.isabs('c:\\folder\\folder'))
    print(os.path.relpath('c:\\folder\\folder\\spam.png', 'c:\\folder1'))
    print(os.path.basename('c:\\folder\\folder\\spam.png'))
    print(os.path.dirname('c:\\folder\\folder\\spam.png'))
    print(os.path.exists('c:\\windows\\system32\\calc.exe'))
    print(os.path.isfile('c:\\windows\\system32\\calc.exe'))
    print(os.path.isdir('c:\\windows\\system32\\calc.exe'))
    print(os.path.getsize('c:\\windows\\system32\\calc.exe'))
    print(os.listdir(r'C:\Projects\MyPythonScripts'))
    # os.makedirs('c:\\delicious\\walnut\\waffles')


def total_size(path):
    totalSize = 0
    for filename in os.listdir(path):
        join = os.path.join(path, filename)
        if not os.path.isfile(join):
            continue
        totalSize = totalSize + os.path.getsize(join)
    return totalSize


def read_file(path):
    hello_file = open(path)
    text = hello_file.read()
    # print(hello_file.readlines())
    hello_file.close()
    return text


# fundamentals()
# read_file()

def write_file():
    hello_file = open(r'C:\Users\as950\Hello2.txt', 'w')
    hello_file.write('Hello!!!!!!!!!!!!!!\n')
    hello_file.close()


# write_file()


def write_file2():
    bacon_file = open('bacon.txt', 'w')
    bacon_file.write('Bacon is not a vegetable.')
    bacon_file.close()


def write_shelve():
    global shelfFile
    shelfFile = shelve.open('mydata')
    shelfFile['cats'] = ['Zophie', 'Pooka', 'Fat-tail', 'Cleo']
    shelfFile.close()


# write_file2()
# write_shelve()

def read_shelve():
    global shelfFile
    shelfFile = shelve.open('mydata')
    print(shelfFile['cats'])
    shelfFile.close()


# read_shelve()


def shutil_fundamentals():
    # shutil.copy('C:\Projects\MyPythonScripts\movies.csv', r'c:\Users\as950\Documents\movies.csv')
    # shutil.copytree(r'c:\delicious', r'c:\Users\as950\Documents\delicious')
    # shutil.move(r'c:\Users\as950\Hello.txt', r'c:\Users\as950\Documents\delicious')
    # shutil.rmtree(r'c:\delicious')
    pass


# shutil_fundamentals()
# os.rmdir(r'c:\Users\as950\Documents\delicious')

def unlink_method():
    os.chdir('C:\\Users\\as950\\Desktop')
    for filename in os.listdir():
        if filename.endswith('.rxt'):
            os.unlink(filename)
            print(filename)


unlink_method()


def send2trash_method():
    os.chdir('C:\\Users\\as950\\Desktop')
    for filename in os.listdir():
        if filename.endswith('.rxt'):
            send2trash.send2trash(filename)
            print(filename)


# send2trash_method()

def walk_method():
    for folder_name, sub_folders, file_names in os.walk('c:\\delicious'):
        print('The folder is ' + folder_name)
        print('The sub_folders in ' + folder_name + 'are ' + str(sub_folders))
        print('The file_names in ' + folder_name + 'are ' + str(file_names))
        print()
        for sub_folder in sub_folders:
            if 'fish' in sub_folder:
                os.rmdir(sub_folder)
        for file in file_names:
            if file.endswith('.py'):
                shutil.copy(os.join(folder_name, file), os.join(folder_name, file + '.backup'))


# walk_method()

# Web Scraping

def mapit():
    if len(sys.argv) > 1:
        address = ' '.join(sys.argv[1:])
    else:
        address = pyperclip.paste()
    webbrowser.open(f'https://www.google.com.eg/maps/place/{address}')


# mapit()

def sel_method():
    browser = webdriver.Chrome(executable_path='C:\\chromedriver.exe')
    browser.get('https://esco.egybest.network/episode/ozark-season-1-ep-1/?refresh=5')
    # elem = browser.find_element_by_css_selector('body')
    # elem.click()
    # elems = browser.find_element_by_css_selector('p')
    # print(len(elems))
    # searchElem = browser.find_element_by_css_selector('#head > div > div.topsrch.suba_rel.nohide.td.vam > form > input.q.autoComplete')
    # searchElem.send_keys('The Dark Knight')
    # searchElem.submit()
    # browser.back()
    # browser.forward()
    # browser.refresh()
    # browser.quit()


# EXCEL

def read_excel():
    global cell
    os.chdir('c:\\Users\\as950\\Documents')
    workbook = openpyxl.load_workbook('PLStandings.xlsx')
    print(type(workbook))
    sheet = workbook.get_sheet_by_name('Sheet1')
    print(workbook.get_sheet_names())
    cell = sheet['B2']
    print(cell.value)
    print(sheet.cell(row=1, column=2))
    for i in range(1, 8):
        print(i, sheet.cell(row=i, column=2).value)


def editing_excel():
    wb = openpyxl.Workbook()
    print(wb.get_sheet_names())
    sheet = wb.get_sheet_by_name('Sheet')
    sheet['A1'] = 'Man City'
    sheet['A2'] = 'Liverpool'
    os.chdir('c:\\Users\\as950\\Documents')
    wb.save('example.xlsx')
    sheet2 = wb.create_sheet()
    sheet2.title = 'My New Sheet Name'
    print(wb.get_sheet_names())
    wb.save('example2.xlsx')
    wb.create_sheet(index=0, title='My other Sheet')
    wb.save('example3.xlsx')


# PDF
def pdf_method():
    global title
    title = 'Al Sweigart - Automate the Boring Stuff with Python, 2nd Edition_ Practical Programming for Total Beginners-No Starch Press (2019).pdf'
    os.chdir(r'C:\Users\as950\Documents\Books')
    pdf_file = open(title, 'rb')
    reader = PyPDF2.PdfFileReader(pdf_file)
    print(reader.numPages)
    page = reader.getPage(2)
    print(page.extractText())
    pdf_file.close()


# WORD

def word_fundamentals():
    global p
    d = docx.Document('c:\\Users\\as950\\Downloads\\demo.docx')
    print(len(d.paragraphs))
    print(d.paragraphs[2].text)
    p = d.paragraphs[1]
    print(len(p.runs))
    print(p.runs[2].text)
    print(p.runs[2].bold)
    p.runs[3].underline = True
    p.runs[3].text = 'Isabelle'
    d.save('c:\\Users\\as950\\Downloads\\demo2.docx')
    print(p.style)
    p.style = 'Title'
    d.save('c:\\Users\\as950\\Downloads\\demo3.docx')
    d = docx.Document()
    d.add_paragraph('Hello this is Isabelle')
    d.add_paragraph('It\'s me again')
    d.save('c:\\Users\\as950\\Downloads\\demo4.docx')
    p = d.paragraphs[0]
    p.add_run('This is new run')
    print(p.runs)
    p.runs[1].bold = True
    d.save('c:\\Users\\as950\\Downloads\\demo4.docx')


def getText(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def send_email(address, Subject, Body):
    conn = smtplib.SMTP('smtp.gmail.com', 587)
    conn.ehlo()
    conn.starttls()
    print(conn.login('as9506683@gmail.com', read_file(r'C:\Projects\gmail.txt')))
    conn.sendmail('as9506683@gmail.com', address,
                  f'Subject: {Subject}\n\n{Body}\n-Salem')
    conn.quit()


def check_email():
    conn = imapclient.IMAPClient('imap.gmail.com', ssl=True)
    print(conn.login('as9506683@gmail.com', read_file(r'C:\Projects\gmail.txt')))
    print(conn.select_folder('INBOX', readonly=True))
    UIDs = conn.search(['SINCE', '28-Apr-2022'])
    print(UIDs)
    for msgid, data in conn.fetch(UIDs, ['ENVELOPE']).items():
        envelope = data[b'ENVELOPE']
        print('ID #%d: "%s" received %s' % (msgid, envelope.subject.decode(), envelope.date))
    raw_message = conn.fetch([2973], ['BODY[]', 'FLAGS'])
    message = pyzmail.PzMessage.factory(raw_message[2973][b'BODY[]'])
    print(message.get_subject())
    print(message.get_addresses('from'))
    print(message.get_addresses('bcc'))
    print(message.get_addresses('bcc'))
    print(message.text_part)
    print(message.html_part)
    print(message.text_part.get_payload().decode('UTF-8'))
    print(message)
    print(conn.list_folders())
    conn.select_folder('INBOX', readonly=False)
    UIDs = conn.search(['ON', '24-Apr-2022'])
    print(UIDs)
    # conn.delete_messages([2930])


# Mouse
def mouse_method():
    print(pyautogui.size())
    width, height = pyautogui.size()
    print(pyautogui.position())
    # pyautogui.moveTo(10, 10)
    # pyautogui.moveTo(10, 200, duration = 1.5)
    # pyautogui.moveTo(10, 200, duration = 1.5)
    # pyautogui.moveRel(200, 0, duration=2)
    # pyautogui.click(1352, 16)
    # pyautogui.rightClick(1320, 339)
    # pyautogui.displayMousePosition()


# Keyboard
def keyboard_method():
    # pyautogui.click(100, 100)
    # pyautogui.typewrite(['a', 'b', 'left', 'left', 'X', 'Y'], interval=0.2)
    print(pyautogui.KEYBOARD_KEYS)
    pyautogui.hotkey('ctrl', 'alt', 'l')


# Images
def Image_method():
    # pyautogui.screenshot('c:\\Users\\as950\\Documents\\screenshot_example.png')
    loc = pyautogui.locateCenterOnScreen(r'C:\Projects\key8.png')
    print(loc)
    pyautogui.moveTo(loc, duration=1)
    pyautogui.click(loc)


